#!/usr/bin/env python3
import sys, os, json, hashlib, random, base64, datetime
from cryptography.fernet import Fernet

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QPushButton, QVBoxLayout, QWidget, QDialog,
    QLabel, QLineEdit, QComboBox, QListWidget, QHBoxLayout, QMessageBox,
    QTableWidget, QTableWidgetItem, QCheckBox, QHeaderView, QAbstractItemView,
    QButtonGroup, QRadioButton, QFileDialog, QTextEdit, QMenuBar, QAction
)
from PyQt5.QtCore import Qt
<<<<<<< HEAD
from PyQt5.QtGui import QPixmap, QFont
=======
from PyQt5.QtGui import QPixmap
from docx import Document
>>>>>>> 69d2547181a4d6835c93cd5e3272451c2fb82e53

# ===================== Универсальные стили =====================
label_style = """
    QLabel {
        font-family: "Segoe UI", Arial, sans-serif;
        font-size: 20pt;
        color: #2b2b2b;
        background: #fff;
        padding: 5px;
    }
"""

combobox_style = """
    QComboBox {
        font-family: "Segoe UI", Arial, sans-serif;
        font-size: 12pt;
        padding: 5px;
        border: 2px solid #000;
        border-radius: 5px;
        background-color: #ffffff;
        color: #000;
    }
    QComboBox::drop-down {
        subcontrol-origin: padding;
        subcontrol-position: top right;
        width: 25px;
        border-radius: 5px;
        border-left: 2px solid #000;
    }
    QComboBox QAbstractItemView {
        font-family: "Segoe UI", Arial, sans-serif;
        font-size: 12pt;
        background-color: #ffffff;
        border: 2px solid #000;
        border-radius: 5px;
        selection-background-color: #d9d9d9;
    }
"""

tablewidget_style = """
    QTableWidget {
        font-family: "Segoe UI", Arial, sans-serif;
        font-size: 11pt;
        color: #2b2b2b;
        background-color: #ffffff;
        gridline-color: #aeaeae;
        border: 1px solid #aeaeae;
    }
    QHeaderView::section {
        background-color: #f7f7f7;
        padding: 5px;
        border: 1px solid #aeaeae;
        font-size: 12pt;
        font-weight: bold;
    }
"""

# Стиль для поля ввода
password_edit_style = """
    QLineEdit {
        padding: 10px;
        font-size: 14px;
        border: 2px solid #ccc;
        border-radius: 5px;
        background: #fff;
        color: #000;
    }
    QLineEdit:focus {
        border: 2px solid #6a9eda;
        background: #fff;
    }
"""

button_style = """
    QPushButton {
        background-color: #fff;
        font-size: 20px;
        color: #000000;
        border: 2px solid;
        padding: 10px;
        border-radius: 5px;
    }
    QPushButton:hover {
        background-color: #aeaeae;
    }
"""

button_not_enable = """
    QPushButton {
        font-size: 20px;
        color: #000000;
        border: 2px solid;
        padding: 10px;
        border-radius: 5px;
        background-color: #aeaeae;
    }
"""

q_message_box = """
    QMessageBox {
        background-color: #fff;
        font-size: 14px;
    }
    QMessageBox QLabel {
        color: #000; /* Цвет текста */
        font-weight: bold;
    }
    QMessageBox QPushButton {
        background-color: #fff;
        font-size: 20px;
        color: #000000;
        border: 2px solid;
        padding: 10px;
        border-radius: 5px;
    }
    QMessageBox QPushButton:hover {
        background-color: #aeaeae; /* Цвет кнопки при наведении */
    }
"""

# ===================== Функции хэширования пароля =====================
def hash_password(password, salt="some_salt"):
    return hashlib.sha256((salt+password).encode()).hexdigest()

def check_password(input_password, stored_hash, salt="some_salt"):
    return hash_password(input_password, salt) == stored_hash

# ===================== Менеджер данных =====================
ENCRYPTION_KEY = b'k_Nm-GI6e-d9U8i8JSGpWn9VZNU3ZQq2gVJXZ32XZxg='

class DataManager:
    def __init__(self, filename="data.enc"):
        self.filename = filename
        self.key = ENCRYPTION_KEY
        self.fernet = Fernet(self.key)
        self.data = self.load_data()

    def load_data(self):
        if not os.path.exists(self.filename):
            default_password = "admin"
            hashed = hash_password(default_password)
            data = {"password": hashed, "tests": [], "results_path": os.path.abspath("results"), "theme": "light"}
            self.save_data(data)
            if not os.path.exists(data["results_path"]):
                os.makedirs(data["results_path"])
            return data
        else:
            with open(self.filename, "rb") as f:
                encrypted_data = f.read()
            try:
                decrypted_data = self.fernet.decrypt(encrypted_data)
                data = json.loads(decrypted_data.decode())
                if "results_path" not in data:
                    data["results_path"] = os.path.abspath("results")
                if "theme" not in data:
                    data["theme"] = "light"
                if not os.path.exists(data["results_path"]):
                    os.makedirs(data["results_path"])
                return data
            except Exception as e:
                print("Ошибка расшифровки данных:", e)
                return {"password": "", "tests": [], "results_path": os.path.abspath("results"), "theme": "light"}

    def save_data(self, data):
        data_json = json.dumps(data, ensure_ascii=False).encode()
        encrypted_data = self.fernet.encrypt(data_json)
        with open(self.filename, "wb") as f:
            f.write(encrypted_data)
        self.data = data

data_manager = DataManager()

# ===================== Менеджер результатов тестов =====================
class ResultManager:
    def __init__(self, results_path):
        self.results_path = results_path
        self.key = ENCRYPTION_KEY
        self.fernet = Fernet(self.key)
        
    def save_result(self, result_data):
        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f")
        filename = os.path.join(self.results_path, f"TEST-{timestamp}.enc")
        data_json = json.dumps(result_data, ensure_ascii=False, indent=2).encode()
        encrypted_data = self.fernet.encrypt(data_json)
        with open(filename, "wb") as f:
            f.write(encrypted_data)
    
    def load_result(self, filepath):
        with open(filepath, "rb") as f:
            encrypted_data = f.read()
        try:
            decrypted_data = self.fernet.decrypt(encrypted_data)
            return json.loads(decrypted_data.decode())
        except Exception as e:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(None, "Ошибка", f"Не удалось расшифровать файл:\n{e}")
            return None

# ===================== Диалоги =====================
class LoginDialog(QDialog):
    def __init__(self, parent=None):
        super(LoginDialog, self).__init__(parent)
        self.setWindowTitle("Вход в режим редактирования")
        layout = QVBoxLayout()
        self.label = QLabel("Введите пароль:")
        self.label.setStyleSheet(label_style)
        self.password_edit = QLineEdit()
        self.password_edit.setStyleSheet(password_edit_style)
        self.password_edit.setEchoMode(QLineEdit.Password)
        self.ok_button = QPushButton("ОК")
        self.ok_button.setStyleSheet(button_style)
        self.ok_button.clicked.connect(self.check_password)
        layout.addWidget(self.label)
        layout.addWidget(self.password_edit)
        layout.addWidget(self.ok_button)
        self.setLayout(layout)
        self.accepted = False

    def check_password(self):
        input_password = self.password_edit.text()
        stored_hash = data_manager.data.get("password", "")
        if check_password(input_password, stored_hash):
            self.accepted = True
            self.accept()
        else:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Неверный пароль!")

class ChangePasswordDialog(QDialog):
    def __init__(self, parent=None):
        super(ChangePasswordDialog, self).__init__(parent)
        self.setWindowTitle("Сменить пароль")
        layout = QVBoxLayout()
        self.current_label = QLabel("Текущий пароль:")
        self.current_label.setStyleSheet(label_style)
        self.current_edit = QLineEdit()
        self.current_edit.setStyleSheet(password_edit_style)
        self.current_edit.setEchoMode(QLineEdit.Password)
        self.new_label = QLabel("Новый пароль:")
        self.new_label.setStyleSheet(label_style)
        self.new_edit = QLineEdit()
        self.new_edit.setStyleSheet(password_edit_style)
        self.new_edit.setEchoMode(QLineEdit.Password)
        self.confirm_label = QLabel("Подтвердите новый пароль:")
        self.confirm_label.setStyleSheet(label_style)
        self.confirm_edit = QLineEdit()
        self.confirm_edit.setStyleSheet(password_edit_style)
        self.confirm_edit.setEchoMode(QLineEdit.Password)
        self.ok_button = QPushButton("Сменить пароль")
        self.ok_button.setStyleSheet("padding: 10px;")
        self.ok_button.clicked.connect(self.change_password)
        layout.addWidget(self.current_label)
        layout.addWidget(self.current_edit)
        layout.addWidget(self.new_label)
        layout.addWidget(self.new_edit)
        layout.addWidget(self.confirm_label)
        layout.addWidget(self.confirm_edit)
        layout.addWidget(self.ok_button)
        self.setLayout(layout)

    def change_password(self):
        current = self.current_edit.text()
        new = self.new_edit.text()
        confirm = self.confirm_edit.text()
        stored_hash = data_manager.data.get("password", "")
        if not check_password(current, stored_hash):
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Неверный пароль!")
            return
        if new != confirm or not new:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Новый пароль и подтверждение не совпадают или пусты!")
            return
        data_manager.data["password"] = hash_password(new)
        data_manager.save_data(data_manager.data)
        msg_box = QMessageBox()
        msg_box.setStyleSheet(q_message_box)
        msg_box.information(self, "Успех", "Пароль успешно изменён!")
        self.accept()

class StudentLoginDialog(QDialog):
    def __init__(self, parent=None):
        super(StudentLoginDialog, self).__init__(parent)
        self.setWindowTitle("Вход студента")
        layout = QVBoxLayout()
        self.label = QLabel("Введите ваше имя:")
        self.label.setStyleSheet(label_style)
        self.name_edit = QLineEdit()
        self.name_edit.setStyleSheet(password_edit_style)
        self.ok_button = QPushButton("Начать тест")
        self.ok_button.setStyleSheet("padding: 10px;")
        self.ok_button.clicked.connect(self.accept)
        self.ok_button.setStyleSheet(button_style)
        layout.addWidget(self.label)
        layout.addWidget(self.name_edit)
        layout.addWidget(self.ok_button)
        self.setLayout(layout)
    
    def get_name(self):
        return self.name_edit.text().strip()

class TestSelectionDialog(QDialog):
    def __init__(self, parent=None):
        super(TestSelectionDialog, self).__init__(parent)
        self.setWindowTitle("Выбор теста")
        layout = QVBoxLayout()
        self.test_list = QListWidget()
        self.test_list.setStyleSheet("""
            QListWidget {
                font-family: "Segoe UI", Arial, sans-serif;
                font-size: 12pt;
                padding: 5px;
                background-color: #ffffff;
                border: 1px solid #d9d9d9;
                border-radius: 5px;
            }
        """)
        self.load_tests()
        self.select_button = QPushButton("Выбрать")
        self.select_button.setStyleSheet(button_style)
        self.select_button.clicked.connect(self.select_test)
        layout.addWidget(self.test_list)
        layout.addWidget(self.select_button)
        self.setLayout(layout)
        self.selected_test = None

    def load_tests(self):
        self.test_list.clear()
        tests = data_manager.data.get("tests", [])
        for test in tests:
            self.test_list.addItem(test.get("topic", "Без темы"))

    def select_test(self):
        current_row = self.test_list.currentRow()
        if current_row >= 0:
            self.selected_test = data_manager.data["tests"][current_row]
            self.accept()
        else:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Выберите тест!")

# ===================== Окно проведения теста =====================
class TestWindow(QDialog):
    def __init__(self, test_data, student_name, parent=None):
        super(TestWindow, self).__init__(parent)
        self.setWindowTitle(f"Тест: {test_data.get('topic','')}")
        self.setMinimumSize(800, 600)
        self.test_data = test_data
        self.student_name = student_name
        self.questions = test_data.get("questions", [])
        self.total_questions = len(self.questions)
        self.question_order = []
        for q in self.questions:
            order = list(range(len(q.get("answers", []))))
            random.shuffle(order)
            self.question_order.append(order)
        self.user_answers = [None] * self.total_questions
        self.current_index = 0

        self.layout = QVBoxLayout()
        self.nav_layout = QHBoxLayout()
        self.prev_button = QPushButton("Предыдущий")
        self.prev_button.setStyleSheet(button_style)
        self.prev_button.clicked.connect(self.go_prev)
        self.nav_combo = QComboBox()
        self.nav_combo.setStyleSheet(combobox_style)
        self.nav_combo.addItems([f"Вопрос {i+1}" for i in range(self.total_questions)])
        self.nav_combo.currentIndexChanged.connect(self.on_nav_change)
        self.next_button = QPushButton("Следующий")
        self.next_button.setStyleSheet(button_style)
        self.next_button.clicked.connect(self.go_next)
        self.finish_button = QPushButton("Завершить тест")
        self.finish_button.setStyleSheet(button_style)
        self.finish_button.clicked.connect(self.finish_test_clicked)
        self.nav_layout.addWidget(self.prev_button)
        self.nav_layout.addWidget(self.nav_combo)
        self.nav_layout.addWidget(self.next_button)
        self.nav_layout.addWidget(self.finish_button)
        self.layout.addLayout(self.nav_layout)

        self.question_image_label = QLabel()
        self.layout.addWidget(self.question_image_label)
        self.question_label = QLabel()
        self.question_label.setStyleSheet("""
            QLabel {
                font-size: 16pt;
                font-weight: bold;
                padding: 10px;
                border: 2px solid #aeaeae;
                border-radius: 8px;
                background-color: #f7f7f7;
            }
        """)
        self.layout.addWidget(self.question_label)

        self.answers_widget = QWidget()
        self.answers_layout = QVBoxLayout()
        self.answers_widget.setLayout(self.answers_layout)
        self.layout.addWidget(self.answers_widget)

        self.setLayout(self.layout)
        self.current_answer_widgets = []
        self.show_question()

    def save_current_answer(self):
        if self.total_questions == 0:
            return
        question = self.questions[self.current_index]
        if question.get("type") == "single":
            selected = None
            for rb in self.current_answer_widgets:
                if rb.isChecked():
                    selected = rb.original_index
                    break
            self.user_answers[self.current_index] = selected
        elif question.get("type") == "multiple":
            total = len(question.get("answers", []))
            selection = [False] * total
            for cb in self.current_answer_widgets:
                if cb.isChecked():
                    selection[cb.original_index] = True
            self.user_answers[self.current_index] = selection

    def restore_answer(self):
        ans = self.user_answers[self.current_index]
        if ans is None:
            return
        question = self.questions[self.current_index]
        if question.get("type") == "single":
            for rb in self.current_answer_widgets:
                if rb.original_index == ans:
                    rb.setChecked(True)
        elif question.get("type") == "multiple":
            for cb in self.current_answer_widgets:
                if ans[cb.original_index]:
                    cb.setChecked(True)

    def show_question(self):
        if self.current_index < self.total_questions:
            question = self.questions[self.current_index]
            order = self.question_order[self.current_index]
            if question.get("image"):
                data = base64.b64decode(question["image"])
                pixmap = QPixmap()
                pixmap.loadFromData(data)
                self.question_image_label.setPixmap(pixmap.scaledToWidth(300, Qt.SmoothTransformation))
            else:
                self.question_image_label.clear()
            self.question_label.setText(f"Вопрос {self.current_index+1}: {question.get('question','')}")
            self.question_label.setStyleSheet(
                """
                    QWidget {
                            font-family: "Segoe UI", Arial, sans-serif;
                            font-size: 20pt;
                            border-radius: 5px;
                            margin-bottom: 5px;
                            background-color: #ffffff;
                        }
                """
            )
            for i in reversed(range(self.answers_layout.count())):
                widget = self.answers_layout.itemAt(i).widget()
                if widget:
                    widget.setParent(None)
            self.current_answer_widgets = []
            answers = question.get("answers", [])
            if question.get("type") == "single":
                self.button_group = QButtonGroup(self)
                for orig_idx in order:
                    answer = answers[orig_idx]
                    container = QWidget()
                    container_layout = QHBoxLayout()
                    container_layout.setContentsMargins(5, 5, 5, 5)
                    rb = QRadioButton(answer.get("text", ""))
                    rb.setStyleSheet("""
                        QRadioButton {
                            background: #fff;
                            font-size: 12pt;
                            padding: 5px;
                        }
                    """)
                    rb.original_index = orig_idx
                    rb.correct = answer.get("correct", False)
                    self.button_group.addButton(rb)
                    container_layout.addWidget(rb)
                    if answer.get("image"):
                        data = base64.b64decode(answer["image"])
                        pixmap = QPixmap()
                        pixmap.loadFromData(data)
                        img_label = QLabel()
                        img_label.setPixmap(pixmap.scaledToWidth(100, Qt.SmoothTransformation))
                        container_layout.addWidget(img_label)
                    container.setLayout(container_layout)
                    container.setStyleSheet("""
                        QWidget {
                            border-radius: 5px;
                            margin-bottom: 5px;
                            background-color: #ffffff;
                        }
                    """)
                    self.answers_layout.addWidget(container)
                    self.current_answer_widgets.append(rb)
            elif question.get("type") == "multiple":
                for orig_idx in order:
                    answer = answers[orig_idx]
                    container = QWidget()
                    container_layout = QHBoxLayout()
                    container_layout.setContentsMargins(5, 5, 5, 5)
                    cb = QCheckBox(answer.get("text", ""))
                    cb.setStyleSheet("""
                        QCheckBox {
                            font-size: 12pt;
                            padding: 5px;
                        }
                    """)
                    cb.original_index = orig_idx
                    cb.correct = answer.get("correct", False)
                    cb.penalty = float(answer.get("penalty", 0))
                    container_layout.addWidget(cb)
                    if answer.get("image"):
                        data = base64.b64decode(answer["image"])
                        pixmap = QPixmap()
                        pixmap.loadFromData(data)
                        img_label = QLabel()
                        img_label.setPixmap(pixmap.scaledToWidth(100, Qt.SmoothTransformation))
                        container_layout.addWidget(img_label)
                    container.setLayout(container_layout)
                    container.setStyleSheet("""
                        QWidget {
                            border-radius: 5px;
                            margin-bottom: 5px;
                            background-color: #ffffff;
                        }
                    """)
                    self.answers_layout.addWidget(container)
                    self.current_answer_widgets.append(cb)
            self.restore_answer()
            self.update_nav_controls()
        else:
            self.finish_test()

    def update_nav_controls(self):
        self.nav_combo.blockSignals(True)
        self.nav_combo.setCurrentIndex(self.current_index)
        self.nav_combo.blockSignals(False)
        self.prev_button.setEnabled(self.current_index > 0)
        self.next_button.setEnabled(self.current_index < self.total_questions - 1)
        if (self.current_index >= self.total_questions - 1):
            self.next_button.setStyleSheet(button_not_enable)
        else: 
            self.next_button.setStyleSheet(button_style)
        if (self.current_index == 0):
            self.prev_button.setStyleSheet(button_not_enable)
        else:
            self.prev_button.setStyleSheet(button_style)

    def on_nav_change(self, index):
        if index == self.current_index:
            return
        self.save_current_answer()
        self.current_index = index
        self.show_question()

    def go_prev(self):
        if self.current_index > 0:
            self.save_current_answer()
            self.current_index -= 1
            self.show_question()

    def go_next(self):
        if self.current_index < self.total_questions - 1:
            self.save_current_answer()
            self.current_index += 1
            self.show_question()

    def finish_test_clicked(self):
        self.save_current_answer()
        missing = []
        for i, question in enumerate(self.questions):
            ans = self.user_answers[i]
            if question.get("type") == "single":
                if ans is None:
                    missing.append(i+1)
            elif question.get("type") == "multiple":
                if ans is None or (isinstance(ans, list) and not any(ans)):
                    missing.append(i+1)
        if missing:
            msg = "Следующие вопросы не имеют ни одного ответа:\n" + ", ".join(map(str, missing))
            msg += "\nЗаполните их!"
            reply = QMessageBox.question(self, "Внимание", msg, QMessageBox.Ok)
            if reply == QMessageBox.Ok:
                return
        self.finish_test()

    def save_results_to_docx(self, total_score, percent):
        doc = Document()
        
        # Добавляем заголовок
        doc.add_heading('Результаты теста', 0)
        
        # Добавляем информацию о тесте
        doc.add_paragraph(f"Тема теста: {self.test_data.get('topic', '')}")
        doc.add_paragraph(f"Количество вопросов: {self.total_questions}")
        doc.add_paragraph(f"Ваш суммарный балл: {total_score:.2f} из {self.total_questions}")
        doc.add_paragraph(f"Успешность: {percent:.0f}%")
        
        # Добавляем раздел с ответами
        doc.add_heading('Ответы на вопросы', level=1)
        
        for i, question in enumerate(self.questions):
            doc.add_paragraph(f"Вопрос {i+1}: {question.get('question', '')}")
            
            user_answer = self.user_answers[i]
            answers = question.get("answers", [])
            
            if question.get("type") == "single":
                if user_answer is not None:
                    doc.add_paragraph(f"Ваш ответ: {answers[user_answer].get('text', '')}")
                else:
                    doc.add_paragraph("Ваш ответ: Нет ответа")
            elif question.get("type") == "multiple":
                if user_answer is not None:
                    selected_answers = [answers[idx].get('text', '') for idx, selected in enumerate(user_answer) if selected]
                    doc.add_paragraph(f"Ваши ответы: {', '.join(selected_answers) if selected_answers else 'Нет ответа'}")
                else:
                    doc.add_paragraph("Ваши ответы: Нет ответа")
            
            correct_answers = [ans.get('text', '') for ans in answers if ans.get('correct', False)]
            doc.add_paragraph(f"Правильные ответы: {', '.join(correct_answers)}")
            doc.add_paragraph()  # Пустая строка для разделения вопросов
        
        # Сохраняем документ
        filename, _ = QFileDialog.getSaveFileName(self, "Сохранить результаты теста", "", "Word Documents (*.docx)")
        if filename:
            if not filename.endswith('.docx'):
                filename += '.docx'
            doc.save(filename)
            QMessageBox.information(self, "Успех", f"Результаты сохранены в файл {filename}")


    def finish_test(self):
        total_score = 0
        detailed_results = []
        for i, question in enumerate(self.questions):
            q_text = question.get("question", "")
            q_score = 0
            answers = question.get("answers", [])
            if question.get("type") == "single":
                correct_idx = None
                for idx, ans in enumerate(answers):
                    if ans.get("correct", False):
                        correct_idx = idx
                        break
                user_ans = self.user_answers[i]
                q_score = 1 if user_ans == correct_idx else 0
                correct_answer = answers[correct_idx]["text"] if correct_idx is not None and correct_idx < len(answers) else "Нет правильного ответа"
                detailed_results.append({
                    "question": q_text,
                    "correct_answer": correct_answer,
                    "user_answer": answers[user_ans]["text"] if user_ans is not None and user_ans < len(answers) else "Нет ответа",
                    "score": q_score
                })
            elif question.get("type") == "multiple":
                total_correct = sum(1 for ans in answers if ans.get("correct", False))
                correct_weight = 1/total_correct if total_correct else 0
                raw_score = 0
                user_selection = self.user_answers[i] if self.user_answers[i] is not None else [False]*len(answers)
                correct_selected = 0
                for idx, selected in enumerate(user_selection):
                    if selected:
                        if answers[idx].get("correct", False):
                            raw_score += correct_weight
                            correct_selected += 1
                        else:
                            raw_score -= float(answers[idx].get("penalty", 0))
                if correct_selected > 0 and raw_score < correct_weight:
                    raw_score = correct_weight
                q_score = max(0, min(raw_score, 1))
                correct_answer = ", ".join([ans["text"] for ans in answers if ans.get("correct", False)])
                detailed_results.append({
                    "question": q_text,
                    "correct_answer": correct_answer,
                    "user_answer": [answers[idx]["text"] for idx, selected in enumerate(user_selection) if selected],
                    "score": q_score
                })
            total_score += q_score
        percent = (total_score / self.total_questions) * 100 if self.total_questions else 0
        msg_box = QMessageBox()
        msg_box.setStyleSheet(q_message_box)
        msg_box.information(self, "Результаты", 
            f"Ваш суммарный балл: {total_score:.2f} из {self.total_questions}\n"
            f"Успешность: {percent:.0f}%")
<<<<<<< HEAD
        result_data = {
            "student": self.student_name,
            "timestamp": datetime.datetime.now().isoformat(),
            "test_topic": self.test_data.get("topic", ""),
            "results": detailed_results,
            "total_score": total_score,
            "total_questions": self.total_questions,
            "percent": percent
        }
        results_path = data_manager.data.get("results_path", os.path.abspath("results"))
        rm = ResultManager(results_path)
        rm.save_result(result_data)
=======
        
        # Сохраняем результаты в документ .docx
        self.save_results_to_docx(total_score, percent)
>>>>>>> 69d2547181a4d6835c93cd5e3272451c2fb82e53
        self.accept()

# ===================== Диалог добавления теста =====================
class AddTestDialog(QDialog):
    def __init__(self, parent=None):
        super(AddTestDialog, self).__init__(parent)
        self.setWindowTitle("Добавить тест")
        layout = QVBoxLayout()
        self.label = QLabel("Название теста:")
        self.label.setStyleSheet(label_style)
        self.topic_edit = QLineEdit()
        self.topic_edit.setStyleSheet(password_edit_style)
        self.ok_button = QPushButton("Добавить")
        self.ok_button.setStyleSheet(button_style)
        self.ok_button.clicked.connect(self.add_test)
        layout.addWidget(self.label)
        layout.addWidget(self.topic_edit)
        layout.addWidget(self.ok_button)
        self.setLayout(layout)
        self.topic = None

    def add_test(self):
        topic = self.topic_edit.text().strip()
        if topic:
            self.topic = topic
            self.accept()
        else:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Введите название теста!")

# ===================== Диалог добавления/редактирования вопроса =====================
class AddQuestionDialog(QDialog):
    def __init__(self, question_data=None, parent=None):
        super(AddQuestionDialog, self).__init__(parent)
        self.setWindowTitle("Добавить вопрос" if question_data is None else "Редактировать вопрос")
        self.setMinimumSize(700, 500)
        self.question_data = question_data
        self.question_image_data = None  
        layout = QVBoxLayout()

        self.question_label = QLabel("Текст вопроса:")
        self.question_label.setStyleSheet(label_style)
        self.question_edit = QLineEdit()
        self.question_edit.setStyleSheet(password_edit_style)
        layout.addWidget(self.question_label)
        layout.addWidget(self.question_edit)

        self.question_image_button = QPushButton("Прикрепить изображение")
        self.question_image_button.setStyleSheet(button_style)
        self.question_image_button.clicked.connect(self.choose_question_image)
        layout.addWidget(self.question_image_button)

        self.type_label = QLabel("Тип вопроса:")
        self.type_label.setStyleSheet(label_style)
        self.type_combo = QComboBox()
        self.type_combo.setStyleSheet(combobox_style)
        self.type_combo.addItems(["Один правильный", "Несколько правильных"])
        layout.addWidget(self.type_label)
        layout.addWidget(self.type_combo)

        self.answers_table = QTableWidget()
        self.answers_table.setStyleSheet(tablewidget_style)
        self.answers_table.setColumnCount(4)
        self.answers_table.setHorizontalHeaderLabels(["Ответ", "Правильный", "Штраф", "Изображение"])
        self.answers_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.answers_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        layout.addWidget(self.answers_table)

        btn_layout = QHBoxLayout()
        self.add_row_btn = QPushButton("Добавить ответ")
        self.add_row_btn.setStyleSheet(button_style)
        self.add_row_btn.clicked.connect(lambda: self.add_row())
        self.remove_row_btn = QPushButton("Удалить выбранный ответ")
        self.remove_row_btn.setStyleSheet(button_style)
        self.remove_row_btn.clicked.connect(self.remove_row)
        btn_layout.addWidget(self.add_row_btn)
        btn_layout.addWidget(self.remove_row_btn)
        layout.addLayout(btn_layout)

        self.save_btn = QPushButton("Сохранить")
        self.save_btn.setStyleSheet(button_style)
        self.save_btn.clicked.connect(self.save_question)
        self.cancel_btn = QPushButton("Отмена")
        self.cancel_btn.setStyleSheet(button_style)
        self.cancel_btn.clicked.connect(self.reject)
        layout.addWidget(self.save_btn)
        layout.addWidget(self.cancel_btn)

        self.setLayout(layout)
        if question_data:
            self.load_question_data()

    def choose_question_image(self):
        filename, _ = QFileDialog.getOpenFileName(self, "Выберите изображение", "", "Images (*.png *.jpg *.jpeg *.bmp)")
        if filename:
            with open(filename, "rb") as f:
                data = f.read()
            self.question_image_data = base64.b64encode(data).decode("utf-8")
            self.question_image_button.setText("Изменить изображение")

    def choose_answer_image(self, btn):
        filename, _ = QFileDialog.getOpenFileName(self, "Выберите изображение", "", "Images (*.png *.jpg *.jpeg *.bmp)")
        if filename:
            with open(filename, "rb") as f:
                data = f.read()
            btn.imageData = base64.b64encode(data).decode("utf-8")
            btn.setText("Изменить изображение")

    def load_question_data(self):
        self.question_edit.setText(self.question_data.get("question", ""))
        if self.question_data.get("image"):
            self.question_image_data = self.question_data.get("image")
            self.question_image_button.setText("Изменить изображение")
        qtype = self.question_data.get("type", "single")
        self.type_combo.setCurrentIndex(0 if qtype=="single" else 1)
        answers = self.question_data.get("answers", [])
        self.answers_table.setRowCount(0)
        for ans in answers:
            self.add_row(ans.get("text", ""), ans.get("correct", False), ans.get("penalty", 0), ans.get("image", ""))

    def add_row(self, text="", correct=False, penalty=0, answer_image=""):
        row = self.answers_table.rowCount()
        self.answers_table.insertRow(row)
        item = QTableWidgetItem(text)
        self.answers_table.setItem(row, 0, item)
        checkbox = QCheckBox()
        checkbox.setChecked(correct)
        self.answers_table.setCellWidget(row, 1, checkbox)
        penalty_combo = QComboBox()
        penalty_combo.setStyleSheet(combobox_style)
        penalty_values = ["0", "0.2", "0.33", "0.4", "0.5", "0.6", "0.67", "0.8", "1"]
        penalty_combo.addItems(penalty_values)
        try:
            penalty_str = "{:.2f}".format(float(penalty)).rstrip("0").rstrip(".")
            if penalty_str == "":
                penalty_str = "0"
            index = penalty_values.index(penalty_str)
        except ValueError:
            index = 0
        penalty_combo.setCurrentIndex(index)
        self.answers_table.setCellWidget(row, 2, penalty_combo)
        image_button = QPushButton("Прикрепить изображение")
        image_button.setStyleSheet(button_style)
        image_button.clicked.connect(lambda _, btn=image_button: self.choose_answer_image(btn))
        image_button.imageData = answer_image
        if answer_image:
            image_button.setText("Изменить изображение")
        self.answers_table.setCellWidget(row, 3, image_button)

    def remove_row(self):
        indices = self.answers_table.selectionModel().selectedRows()
        for index in sorted(indices, reverse=True):
            self.answers_table.removeRow(index.row())

    def save_question(self):
        question_text = self.question_edit.text().strip()
        if not question_text:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Введите текст вопроса!")
            return
        qtype = "single" if self.type_combo.currentIndex() == 0 else "multiple"
        answers = []
        for row in range(self.answers_table.rowCount()):
            answer_item = self.answers_table.item(row, 0)
            if answer_item:
                answer_text = answer_item.text().strip()
                checkbox = self.answers_table.cellWidget(row, 1)
                correct = checkbox.isChecked() if checkbox else False
                penalty_combo = self.answers_table.cellWidget(row, 2)
                try:
                    penalty_val = float(penalty_combo.currentText())
                except Exception:
                    penalty_val = 0
                image_button = self.answers_table.cellWidget(row, 3)
                answer_image = image_button.imageData if hasattr(image_button, "imageData") else ""
                if answer_text:
                    answers.append({
                        "text": answer_text,
                        "correct": correct,
                        "penalty": penalty_val,
                        "image": answer_image
                    })
        if len(answers) < 2:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Должно быть минимум два ответа!")
            return
        if qtype == "single":
            correct_count = sum(1 for ans in answers if ans["correct"])
            if correct_count != 1:
                msg_box = QMessageBox()
                msg_box.setStyleSheet(q_message_box)
                msg_box.warning(self, "Ошибка", "Для вопроса с одним правильным ответом должен быть ровно один правильный вариант!")
                return
        self.question_result = {
            "question": question_text,
            "type": qtype,
            "answers": answers,
            "image": self.question_image_data if self.question_image_data else ""
        }
        self.accept()

# ===================== Окно редактирования тестов и управления результатами =====================
class EditWindow(QMainWindow):
    def __init__(self, parent=None):
        super(EditWindow, self).__init__(parent)
        self.setWindowTitle("Редактирование тестов")
        self.setMinimumSize(800, 600)
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QHBoxLayout()
        main_widget.setLayout(layout)

        self.tests_list = QListWidget()
        self.tests_list.setStyleSheet("""
            QListWidget {
                font-family: "Segoe UI", Arial, sans-serif;
                font-size: 12pt;
                padding: 5px;
                background-color: #ffffff;
                border: 2px solid #000000;
                border-radius: 5px;
            }
        """)
        self.tests_list.itemClicked.connect(self.load_questions)
        layout.addWidget(self.tests_list)

        edit_block = QWidget()
        edit_block.setObjectName("EditBlock")
        edit_block.setStyleSheet("""
            QWidget#EditBlock {
                border: 2px solid #aeaeae;
                border-radius: 10px;
                background-color: #f7f7f7;
                padding: 10px;
            }
        """)
        edit_layout = QVBoxLayout()
        edit_block.setLayout(edit_layout)

        self.questions_list = QListWidget()
        self.questions_list.setStyleSheet("""
            QListWidget {
                font-family: "Segoe UI", Arial, sans-serif;
                font-size: 12pt;
                padding: 5px;
                background-color: #ffffff;
                border: 2px solid #000000;
                border-radius: 5px;
            }
        """)
        edit_layout.addWidget(self.questions_list)

        btn_layout = QVBoxLayout()
        self.add_test_btn = QPushButton("Добавить тест")
        self.add_test_btn.setStyleSheet(button_style)
        self.add_test_btn.clicked.connect(self.add_test)
        self.delete_test_btn = QPushButton("Удалить тест")
        self.delete_test_btn.setStyleSheet(button_style)
        self.delete_test_btn.clicked.connect(self.delete_test)
        self.add_question_btn = QPushButton("Добавить вопрос")
        self.add_question_btn.setStyleSheet(button_style)
        self.add_question_btn.clicked.connect(self.add_question)
        self.delete_question_btn = QPushButton("Удалить вопрос")
        self.delete_question_btn.setStyleSheet(button_style)
        self.delete_question_btn.clicked.connect(self.delete_question)
        self.edit_question_btn = QPushButton("Редактировать вопрос")
        self.edit_question_btn.setStyleSheet(button_style)
        self.edit_question_btn.clicked.connect(self.edit_question)
        self.change_password_btn = QPushButton("Сменить пароль")
        self.change_password_btn.setStyleSheet(button_style)
        self.change_password_btn.clicked.connect(self.change_password)
        self.set_results_path_btn = QPushButton("Изменить путь сохранения результатов")
        self.set_results_path_btn.setStyleSheet(button_style)
        self.set_results_path_btn.clicked.connect(self.set_results_path)
        self.view_results_btn = QPushButton("Просмотр результатов")
        self.view_results_btn.setStyleSheet(button_style)
        self.view_results_btn.clicked.connect(self.view_results)
        for btn in [self.add_test_btn, self.delete_test_btn, self.add_question_btn,
                    self.delete_question_btn, self.edit_question_btn, self.change_password_btn,
                    self.set_results_path_btn, self.view_results_btn]:
            btn_layout.addWidget(btn)
        btn_layout.addStretch()
        edit_layout.addLayout(btn_layout)

        layout.addWidget(edit_block)

        self.load_tests()

    def load_tests(self):
        self.tests_list.clear()
        tests = data_manager.data.get("tests", [])
        for test in tests:
            self.tests_list.addItem(test.get("topic", "Без темы"))
        self.questions_list.clear()

    def load_questions(self):
        current_row = self.tests_list.currentRow()
        if current_row >= 0:
            test = data_manager.data["tests"][current_row]
            self.questions_list.clear()
            for q in test.get("questions", []):
                self.questions_list.addItem(q.get("question", ""))
    
    def add_test(self):
        dialog = AddTestDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            new_test = {"topic": dialog.topic, "questions": []}
            data_manager.data["tests"].append(new_test)
            data_manager.save_data(data_manager.data)
            self.load_tests()

    def delete_test(self):
        current_row = self.tests_list.currentRow()
        if current_row >= 0:
            confirm = QMessageBox.question(self, "Подтверждение", "Удалить выбранный тест?")
            if confirm == QMessageBox.Yes:
                del data_manager.data["tests"][current_row]
                data_manager.save_data(data_manager.data)
                self.load_tests()

    def add_question(self):
        current_test_index = self.tests_list.currentRow()
        if current_test_index < 0:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Выберите тест!")
            return
        dialog = AddQuestionDialog(None, self)
        if dialog.exec_() == QDialog.Accepted:
            question = dialog.question_result
            data_manager.data["tests"][current_test_index].setdefault("questions", []).append(question)
            data_manager.save_data(data_manager.data)
            self.load_questions()

    def delete_question(self):
        current_test_index = self.tests_list.currentRow()
        current_question_index = self.questions_list.currentRow()
        if current_test_index < 0 or current_question_index < 0:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Выберите вопрос для удаления!")
            return
        confirm = QMessageBox.question(self, "Подтверждение", "Удалить выбранный вопрос?")
        if confirm == QMessageBox.Yes:
            del data_manager.data["tests"][current_test_index]["questions"][current_question_index]
            data_manager.save_data(data_manager.data)
            self.load_questions()

    def edit_question(self):
        current_test_index = self.tests_list.currentRow()
        current_question_index = self.questions_list.currentRow()
        if current_test_index < 0 or current_question_index < 0:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Выберите вопрос для редактирования!")
            return
        question_data = data_manager.data["tests"][current_test_index]["questions"][current_question_index]
        dialog = AddQuestionDialog(question_data, self)
        if dialog.exec_() == QDialog.Accepted:
            data_manager.data["tests"][current_test_index]["questions"][current_question_index] = dialog.question_result
            data_manager.save_data(data_manager.data)
            self.load_questions()

    def change_password(self):
        dialog = ChangePasswordDialog(self)
        dialog.exec_()

    def set_results_path(self):
        directory = QFileDialog.getExistingDirectory(self, "Выберите директорию для сохранения результатов")
        if directory:
            data_manager.data["results_path"] = directory
            data_manager.save_data(data_manager.data)
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Путь сохранения результатов изменён.")

    def view_results(self):
        results_path = data_manager.data.get("results_path", os.path.abspath("results"))
        if not os.path.exists(results_path):
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.warning(self, "Ошибка", "Директория с результатами не существует.")
            return
        files = [f for f in os.listdir(results_path) if f.startswith("TEST-") and f.endswith(".enc")]
        if not files:
            msg_box = QMessageBox()
            msg_box.setStyleSheet(q_message_box)
            msg_box.information(self, "Информация", "Файлы с результатами не найдены.")
            return
        dialog = QDialog(self)
        dialog.setWindowTitle("Просмотр результатов")
        dialog.setMinimumSize(600, 400)
        layout = QVBoxLayout()
        list_widget = QListWidget()
        list_widget.addItems(files)
        layout.addWidget(list_widget)
        result_text = QTextEdit()
        result_text.setReadOnly(True)
        layout.addWidget(result_text)
        btn_layout = QHBoxLayout()
        open_btn = QPushButton("Открыть")
        open_btn.setStyleSheet(button_style)
        close_btn = QPushButton("Закрыть")
        close_btn.setStyleSheet(button_style)
        btn_layout.addWidget(open_btn)
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)
        dialog.setLayout(layout)

        rm = ResultManager(results_path)

        def format_result(result):
            try:
                dt = datetime.datetime.fromisoformat(result.get("timestamp", ""))
                time_str = dt.strftime("%Y.%m.%d %H:%M:%S")
            except Exception:
                time_str = result.get("timestamp", "")
            output = []
            output.append(f'Тестируемый: "{result.get("student", "")}"')
            output.append(f'Время начала теста: {time_str}')
            output.append(f'Тест: "{result.get("test_topic", "")}"\n')
            output.append("Результаты:\n")
            for i, item in enumerate(result.get("results", []), 1):
                output.append(f'Вопрос {i}: "{item.get("question", "")}"')
                output.append(f'Правильный ответ: {item.get("correct_answer", "")}')
                ua = item.get("user_answer", "")
                if isinstance(ua, list):
                    ua = ", ".join(ua)
                output.append(f'Ответ тестируемого: {ua}')
                output.append(f'Баллы: {item.get("score", 0)}\n')
            output.append(f'Итоговый балл: {result.get("total_score", 0)}')
            output.append(f'Всего вопросов: {result.get("total_questions", 0)}')
            output.append(f'Процент прохождения: {result.get("percent", 0)}%')
            return "\n".join(output)

        def open_result():
            selected = list_widget.currentItem()
            if selected:
                filepath = os.path.join(results_path, selected.text())
                result = rm.load_result(filepath)
                if result:
                    formatted = format_result(result)
                    result_text.setText(formatted)
        open_btn.clicked.connect(open_result)
        close_btn.clicked.connect(dialog.accept)
        dialog.exec_()

# ===================== Главное окно =====================
class MainWindow(QMainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        self.setWindowTitle("Приложение для тестирования")
        self.setMinimumSize(800, 600)
        self.create_menu()
        self.setStyleSheet("""
            background: qlineargradient(
                x1: 0, y1: 0, x2: 1, y2: 1,
                stop: 0 #A4D4F9, stop: 1 #005392
            );
        """)
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout()
        central_widget.setLayout(layout)

        title_label = QLabel("Приложение для тестирования")
        title_label.setFont(QFont("Segoe UI", 18, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)

        test_button = QPushButton("Тестирование")
        test_button.setFont(QFont("Segoe UI", 14))
        test_button.setStyleSheet(button_style)
        test_button.clicked.connect(self.start_test)
        layout.addWidget(test_button)

        edit_button = QPushButton("Редактирование")
        edit_button.setFont(QFont("Segoe UI", 14))
        edit_button.setStyleSheet(button_style)
        edit_button.clicked.connect(self.start_edit)
        layout.addWidget(edit_button)

    def create_menu(self):
        menubar = QMenuBar(self)
        self.setMenuBar(menubar)

    def start_test(self):
        login_dialog = StudentLoginDialog(self)
        if login_dialog.exec_() == QDialog.Accepted:
            student_name = login_dialog.get_name()
            if not student_name:
                msg_box = QMessageBox()
                msg_box.setStyleSheet(q_message_box)
                msg_box.warning(self, "Ошибка", "Введите имя!")
                return
            dialog = TestSelectionDialog(self)
            if dialog.exec_() == QDialog.Accepted:
                test_data = dialog.selected_test
                test_window = TestWindow(test_data, student_name, self)
                test_window.exec_()

    def start_edit(self):
        login_dialog = LoginDialog(self)
        if login_dialog.exec_() == QDialog.Accepted and login_dialog.accepted:
            self.edit_window = EditWindow(self)
            self.edit_window.show()

# ===================== Точка входа в приложение =====================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    main_win = MainWindow()
    main_win.show()
    sys.exit(app.exec_())
